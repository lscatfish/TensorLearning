# -*- coding: utf-8 -*-
import re
import os
import tempfile
import zipfile
from copy import deepcopy

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree
import pypandoc

# ========================= 配置 =========================
CHINESE_FONT = '宋体'
WESTERN_FONT = 'Times New Roman'
COLOR_BLACK = RGBColor(0, 0, 0)

TITLE_SIZES = {
    'h1': Pt(22),
    'h2': Pt(16),
    'h3': Pt(14),
    'h4': Pt(12),
}

BODY_SIZE = Pt(12)
QUOTE_SIZE = Pt(10.5)
FIRST_LINE_INDENT = Cm(0.85)

MATH_PLACEHOLDER_RE = re.compile(r'\{\{MATH:(\d+)\}\}')

# 全局公式存储
formulas = []   # (is_display, latex)
omml_cache = {}  # idx -> lxml Element


# ========================= 公式处理 =========================
def extract_formulas(md_text):
    """提取块级和行内公式，替换为占位符"""
    def repl_display(m):
        idx = len(formulas)
        formulas.append((True, m.group(1).strip()))
        return f'\n\n{{{{MATH:{idx}}}}}\n\n'

    text = re.sub(r'\$\$\s*(.*?)\s*\$\$', repl_display, md_text, flags=re.DOTALL)

    def repl_inline(m):
        idx = len(formulas)
        formulas.append((False, m.group(1).strip()))
        return f'{{{{MATH:{idx}}}}}'

    text = re.sub(r'(?<!\$)\$(?!\$)([^\$\n]+?)\$(?!\$)', repl_inline, text)
    return text


def get_omml(idx):
    if idx in omml_cache:
        return omml_cache[idx]

    is_display, latex = formulas[idx]

    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8', dir='.') as f:
        if is_display:
            f.write(f'$${latex}$$\n')
        else:
            f.write(f'${latex}$\n')
        temp_md = f.name

    temp_docx = temp_md.replace('.md', '.docx')
    pypandoc.convert_file(temp_md, 'docx', outputfile=temp_docx)

    doc = Document(temp_docx)
    ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}

    result = None
    for p in doc.paragraphs:
        p_xml = p._p
        omath_para = p_xml.find('.//m:oMathPara', namespaces=ns)
        if omath_para is not None:
            result = deepcopy(omath_para)
            break
        omath = p_xml.find('.//m:oMath', namespaces=ns)
        if omath is not None:
            result = deepcopy(omath)
            break

    os.remove(temp_md)
    os.remove(temp_docx)

    omml_cache[idx] = result
    return result


def replace_math_in_docx(docx_path):
    """在 docx 的 XML 中替换占位符为 OMML"""
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 读取整个 zip 到内存
    with zipfile.ZipFile(docx_path, 'r') as zin:
        items = []
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                root = etree.fromstring(data)

                for p in root.iter(f'{{{W_NS}}}p'):
                    runs = list(p.iter(f'{{{W_NS}}}r'))
                    for r in runs:
                        t_elems = list(r.iter(f'{{{W_NS}}}t'))
                        for t in t_elems:
                            text = t.text or ''
                            m = re.match(r'^\{\{MATH:(\d+)\}\}$', text)
                            if m:
                                idx = int(m.group(1))
                                omath = omml_cache.get(idx)
                                if omath is not None:
                                    r.addprevious(deepcopy(omath))
                                    p.remove(r)
                                break

                data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
            items.append((item, data))

    # 写回原文件
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item, data in items:
            zout.writestr(item, data)


# ========================= 工具函数 =========================
def remove_spaces_between_languages(text):
    if not text:
        return text
    text = re.sub(r'([\u4e00-\u9fff\u3000-\u303f\uff00-\uffef])\s+([\w%+\-<>=^*/()\[\]{}])', r'\1\2', text)
    text = re.sub(r'([\w%+\-<>=^*/()\[\]{}])\s+([\u4e00-\u9fff\u3000-\u303f\uff00-\uffef])', r'\1\2', text)
    text = re.sub(r'\s{2,}', ' ', text)
    text = text.strip()
    return text


def set_run_font(run, cn_font=CHINESE_FONT, en_font=WESTERN_FONT, size=None, bold=False, italic=False, color=COLOR_BLACK):
    run.font.name = en_font
    rFonts = run._element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), cn_font)
    # Chinese punctuation (e.g. “” ‘’) uses hAnsi font; set it to Chinese font
    rFonts.set(qn('w:hAnsi'), cn_font)
    if size:
        run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_text_to_paragraph(paragraph, text, size=None, bold=False, italic=False):
    """添加文本到段落，自动拆分公式占位符"""
    parts = MATH_PLACEHOLDER_RE.split(text)
    for i, part in enumerate(parts):
        if i % 2 == 1:
            idx = int(part)
            run = paragraph.add_run(f'{{{{MATH:{idx}}}}}')
            set_run_font(run, size=size, bold=bold, italic=italic)
        else:
            text_clean = remove_spaces_between_languages(part)
            if text_clean:
                run = paragraph.add_run(text_clean)
                set_run_font(run, size=size, bold=bold, italic=italic)


def process_inline_elements(paragraph, node, size=None, bold=False, italic=False):
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            parts = MATH_PLACEHOLDER_RE.split(text)
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    idx = int(part)
                    run = paragraph.add_run(f'{{{{MATH:{idx}}}}}')
                    set_run_font(run, size=size, bold=bold, italic=italic)
                else:
                    text_clean = remove_spaces_between_languages(part)
                    if text_clean:
                        run = paragraph.add_run(text_clean)
                        set_run_font(run, size=size, bold=bold, italic=italic)
        elif child.name in ('strong', 'b'):
            process_inline_elements(paragraph, child, size=size, bold=True, italic=italic)
        elif child.name in ('em', 'i'):
            process_inline_elements(paragraph, child, size=size, bold=bold, italic=True)
        elif child.name == 'code':
            text = remove_spaces_between_languages(child.get_text())
            if text:
                run = paragraph.add_run(text)
                set_run_font(run, en_font='Courier New', cn_font='Courier New', size=size, bold=bold, italic=italic)
        elif child.name == 'br':
            paragraph.add_run().add_break()
        elif child.name == 'a':
            text = remove_spaces_between_languages(child.get_text())
            if text:
                run = paragraph.add_run(text)
                set_run_font(run, size=size, bold=bold, italic=italic)
                run.font.underline = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        else:
            process_inline_elements(paragraph, child, size=size, bold=bold, italic=italic)


# ========================= 主处理 =========================
def convert_md_to_docx(md_path, docx_path):
    global formulas, omml_cache
    formulas.clear()
    omml_cache.clear()

    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 去掉 SVG 图片引用
    md_text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'[图：\1]', md_text)

    # 提取公式
    md_text = extract_formulas(md_text)

    # 预生成所有 OMML
    print(f"发现 {len(formulas)} 个公式，正在转换...")
    for i in range(len(formulas)):
        elem = get_omml(i)
        if elem is None:
            print(f"  警告: 公式 {i} 转换失败: {formulas[i]}")
    print("公式转换完成")

    # Markdown -> HTML
    html = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'lxml')

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = WESTERN_FONT
    style._element.rPr.rFonts.set(qn('w:eastAsia'), CHINESE_FONT)
    style.font.size = BODY_SIZE
    style.font.color.rgb = COLOR_BLACK

    body = soup.body if soup.body else soup

    for elem in body.children:
        if isinstance(elem, NavigableString):
            if elem.strip():
                p = doc.add_paragraph()
                add_text_to_paragraph(p, elem.strip(), size=BODY_SIZE)
                p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_after = Pt(6)
            continue

        tag = elem.name

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag[1])
            size = TITLE_SIZES.get(tag, Pt(12))
            p = doc.add_paragraph()
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = elem.get_text().strip()
            text = remove_spaces_between_languages(text)
            run = p.add_run(text)
            set_run_font(run, size=size, bold=True)
            p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.25

        elif tag == 'p':
            p = doc.add_paragraph()
            process_inline_elements(p, elem, size=BODY_SIZE)
            p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(0)

        elif tag == 'blockquote':
            p = doc.add_paragraph()
            process_inline_elements(p, elem, size=QUOTE_SIZE, italic=True)
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(6)

        elif tag == 'table':
            thead = elem.find('thead')
            tbody = elem.find('tbody')
            rows = []
            if thead:
                rows += [[th.get_text(strip=True) for th in tr.find_all(['th', 'td'])] for tr in thead.find_all('tr')]
            if tbody:
                rows += [[td.get_text(strip=True) for td in tr.find_all(['th', 'td'])] for tr in tbody.find_all('tr')]
            else:
                rows += [[td.get_text(strip=True) for td in tr.find_all(['th', 'td'])] for tr in elem.find_all('tr')]

            if not rows:
                continue
            col_count = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, row_data in enumerate(rows):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = remove_spaces_between_languages(cell_text)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            set_run_font(run, size=BODY_SIZE, bold=(i == 0))
                if i == 0:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

            doc.add_paragraph()

        elif tag in ('ul', 'ol'):
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph()
                bullet = '• '
                run = p.add_run(bullet)
                set_run_font(run, size=BODY_SIZE)
                process_inline_elements(p, li, size=BODY_SIZE)
                p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_after = Pt(3)

        elif tag == 'hr':
            # 跳过分割线
            continue

        elif tag == 'pre':
            code = elem.find('code')
            text = code.get_text() if code else elem.get_text()
            p = doc.add_paragraph()
            add_text_to_paragraph(p, text, size=Pt(10.5))
            for run in p.runs:
                run.font.name = 'Courier New'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(6)

        elif tag == 'div':
            for child in elem.children:
                if isinstance(child, NavigableString):
                    if child.strip():
                        p = doc.add_paragraph()
                        add_text_to_paragraph(p, child.strip(), size=BODY_SIZE)
                        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
                elif child.name:
                    p = doc.add_paragraph()
                    process_inline_elements(p, child, size=BODY_SIZE)
                    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    doc.save(docx_path)
    print(f"初步 docx 已保存: {docx_path}")

    # 替换公式占位符为 OMML
    replace_math_in_docx(docx_path)
    print(f"公式已嵌入: {docx_path}")


if __name__ == '__main__':
    convert_md_to_docx('report.md', 'report.docx')
